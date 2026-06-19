#!/usr/bin/env python3
"""
Construye un .docx limpio a partir de un archivo de texto del guión VSL.

Formato de entrada esperado (texto plano UTF-8):
- Una línea que empiece con "# "   -> Título del documento (centrado, grande)
- Una línea que empiece con "## "  -> Encabezado de sección (negrita)
- Líneas que empiecen con "[" y terminen con "]"  -> nota de producción / indicación visual (gris, cursiva)
- Cualquier otra línea no vacía     -> párrafo de guión normal
- Líneas vacías separan párrafos

Uso:
    python build_docx.py <entrada.txt> <salida.docx>
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(input_path: str, output_path: str) -> None:
    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # Cuerpo legible para lectura/locución
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(20)
        elif line.startswith("## "):
            doc.add_paragraph()  # espacio antes de sección
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
        elif line.startswith("[") and line.endswith("]"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"OK -> {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python build_docx.py <entrada.txt> <salida.docx>")
        sys.exit(1)
    build(sys.argv[1], sys.argv[2])
